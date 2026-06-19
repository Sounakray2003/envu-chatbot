"""
DOCX Extractor
Extracts content from Word documents (.docx / .doc).

Strategy:
  1. Try Docling first (handles layout, tables, images natively).
  2. Fall back to python-docx if Docling fails or produces no content.

Returns a plain markdown string. No BaseExtractor dependency.
"""

import logging
import os
from pathlib import Path
from typing import Dict, List

from docx import Document

logger = logging.getLogger(__name__)


class DocxExtractor:
    """
    Word Document Extractor: Docling primary + python-docx fallback.

    Docling handles image description natively (do_picture_description=True).
    No separate Vision LLM / image-cache pipeline.
    Returns a markdown string from extract().

    Relevant env var:
        DOCLING_PICTURE_DESCRIPTION  - 'true'/'false' (default: true)
    """

    def __init__(self):
        import os
        self._picture_description = (
            os.getenv("DOCLING_PICTURE_DESCRIPTION", "true").strip().lower() != "false"
        )
        logger.info(
            f"DOCX Extractor initialized "
            f"(Docling + python-docx fallback, "
            f"picture_description={self._picture_description})"
        )

    # =========================================================================
    # Validation (internal)
    # =========================================================================

    def _validate(self, file_path: str):
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File does not exist: {file_path}")
        if path.suffix.lower() not in [".docx", ".doc"]:
            raise ValueError(f"Not a Word document: {path.suffix}")
        if path.stat().st_size == 0:
            raise ValueError("File is empty (0 bytes)")
        try:
            import zipfile
            if path.suffix.lower() == ".docx" and not zipfile.is_zipfile(file_path):
                raise ValueError("File is not a valid DOCX (not a ZIP archive)")
        except ImportError:
            pass
    
    # =========================================================================
    # Public entry point
    # =========================================================================

    def extract(self, file_path: str) -> str:
        """
        Extract a Word document and return Markdown content.

        Tries Docling first; falls back to python-docx if Docling fails
        or returns empty content.

        Returns:
            Markdown string.

        Raises:
            RuntimeError: If both strategies fail.
        """
        logger.info(f"Extracting DOCX: {file_path}")
        self._validate(file_path)

        # Strategy 1: Docling
        try:
            markdown = self._extract_with_docling(file_path)
            if markdown and markdown.strip():
                logger.info("Docling extraction succeeded")
                return markdown
            logger.warning("Docling produced empty content — falling back to python-docx")
        except Exception as e:
            logger.warning(f"Docling failed ({e}) — falling back to python-docx")

        # Strategy 2: python-docx
        try:
            markdown = self._extract_with_python_docx(file_path)
            if markdown and markdown.strip():
                logger.info("python-docx extraction succeeded")
                return markdown
            raise RuntimeError("python-docx produced empty content")
        except Exception as e:
            raise RuntimeError(
                f"Both Docling and python-docx failed for {file_path}: {e}"
            )
    
    # =========================================================================
    # Strategy 1: Docling
    # =========================================================================

    def _extract_with_docling(self, file_path: str) -> str:
        """
        Extract using Docling with picture description, table structure,
        formula/code enrichment enabled.

        Docling handles image descriptions natively — no separate pipeline needed.
        DOCX files are typically not paginated by Docling, so we export the
        whole document at once when no pages are present.
        """
        from docling.document_converter import DocumentConverter
        from docling.datamodel.pipeline_options import PdfPipelineOptions
        from docling.datamodel.base_models import InputFormat
        from docling.document_converter import WordFormatOption
        from docling.datamodel.pipeline_options import WordPipelineOptions

        doc_stem = Path(file_path).stem
        logger.info("Starting Docling DOCX conversion...")

        pipeline_options = WordPipelineOptions()
        pipeline_options.do_picture_description = self._picture_description
        # Force any Docling artifacts to output_files instead of source directories.
        docx_artifact_dir = Path.cwd() / "output_files" / "docx_artifacts" / doc_stem
        docx_artifact_dir.mkdir(parents=True, exist_ok=True)
        try:
            pipeline_options.artifacts_path = str(docx_artifact_dir)
        except AttributeError:
            # Older/newer Docling versions may not expose artifacts_path for Word.
            pass
        

        converter = DocumentConverter(
            format_options={
                InputFormat.DOCX: WordFormatOption(pipeline_options=pipeline_options)
            }
        )

        result = converter.convert(file_path)

        if not hasattr(result, "document"):
            raise RuntimeError("Docling result has no 'document' attribute")

        document = result.document
        total_pages = len(document.pages) if hasattr(document, "pages") else 0
        logger.info(f"Docling detected {total_pages} page(s) in DOCX")

        return self._build_markdown_docling(document, total_pages, doc_stem)

    def _build_markdown_docling(self, document, total_pages: int, doc_stem: str) -> str:
        """
        Export Docling document to markdown.
        Returns only the actual document content - no auto-generated headers, 
        page markers, or metadata. Image descriptions from Docling are included inline.
        """
        # DOCX has no page structure in Docling — export entire document at once
        if total_pages == 0 or not getattr(document, "pages", None):
            logger.info("DOCX has no Docling pages — exporting full document as single unit")
            try:
                full_markdown = document.export_to_markdown(image_mode="embedded")
                if not full_markdown.strip():
                    raise RuntimeError("No content extracted from document")
                return full_markdown
            except Exception as e:
                logger.error(f"Error exporting DOCX markdown: {e}")
                raise RuntimeError(f"Failed to extract DOCX content: {e}")

        # Paginated path (rare for DOCX)
        result_lines = []
        for page_number in sorted(document.pages.keys()):
            try:
                page_markdown = document.export_to_markdown(
                    page_no=page_number,
                    image_mode="embedded",
                )
                if not page_markdown.strip():
                    continue
                result_lines.append(page_markdown)
            except Exception as e:
                logger.warning(f"Error processing page {page_number}: {e}")

        content = "\n".join(result_lines)
        if not content.strip():
            raise RuntimeError("No content extracted from document")
        return content

    # =========================================================================
    # Strategy 2: python-docx fallback
    # =========================================================================

    def _extract_with_python_docx(self, file_path: str) -> str:
        """
        Extract using python-docx (text + tables only).
        Images are not described here — this is a fallback path when Docling fails.
        """
        logger.info("Starting python-docx extraction...")
        doc = Document(file_path)
        num_paragraphs = len(doc.paragraphs)
        num_tables = len(doc.tables)
        logger.info(f"Document: {num_paragraphs} paragraphs, {num_tables} tables")

        doc_stem = Path(file_path).stem
        markdown = self._build_markdown_python_docx(doc, doc_stem)
        logger.info(f"python-docx markdown: {len(markdown):,} characters")
        return markdown

    def _build_markdown_python_docx(self, doc: Document, doc_stem: str) -> str:
        """
        Build markdown from python-docx paragraphs and tables.
        Returns only document content - no auto-generated headers or page markers.
        """
        result_lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    level_str = style.replace("Heading", "").strip()
                    try:
                        level = max(1, min(int(level_str), 6))
                    except ValueError:
                        level = 2
                    result_lines.append(f"{'#' * level} {text}")
                else:
                    result_lines.append(text)

        for table in doc.tables:
            table_md = self._table_to_markdown(table)
            if table_md:
                result_lines.append(table_md)

        content = "\n".join(result_lines)
        if not content.strip():
            raise RuntimeError("No content extracted from document")
        return content

    def _table_to_markdown(self, table) -> str:
        """Convert a python-docx table to a markdown table string."""
        try:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                return ""
            max_cols = max(len(r) for r in rows)
            rows = [r + [""] * (max_cols - len(r)) for r in rows]
            header = "| " + " | ".join(rows[0]) + " |"
            separator = "| " + " | ".join(["---"] * max_cols) + " |"
            body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
            return "\n".join(filter(None, [header, separator, body]))
        except Exception as e:
            logger.debug(f"Table extraction failed: {e}")
            return ""
